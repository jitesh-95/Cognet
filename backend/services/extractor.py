# backend/services/cleaner.py
# For HTML extraction
from readability import Document as ReadabilityDocument

# For Word documents
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import fitz
from io import BytesIO

def extract_main_html(html: str) -> str:
    """
    Extract the main content from HTML.
    Removes navigation, ads, headers/footers.
    """
    doc = ReadabilityDocument(html)
    content_html = doc.summary()
    # Optional: strip HTML tags to get plain text
    soup = BeautifulSoup(content_html, "html.parser")
    return soup.get_text(separator="\n", strip=True)

def extract_text_from_pdf(pdf_file) -> str:
    """
    Extracts text from the uploaded PDF file.
    """
    doc = fitz.open(pdf_file)
    text = ""
    for page in doc:
        text += page.get_text("text")  # Extract raw text
    return text

def extract_text_from_doc(file_bytes: BytesIO) -> str:
    """
    Extract text from a DOC or DOCX file.

    Args:
        file_bytes (BytesIO): File content as bytes

    Returns:
        str: Extracted text
    """
    doc = DocxDocument(file_bytes)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text